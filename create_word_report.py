from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

def set_background_color(paragraph, color):
    """Paragraf arka plan rengini ayarlar"""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)

def add_header(doc):
    """Sayfa başlığı ekler"""
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Hipertansiyon Tahmin Veri Analizi ve Makine Öğrenmesi"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_footer(doc):
    """Sayfa altlığı ekler"""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = f"Sayfa {datetime.now().strftime('%d/%m/%Y')}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def create_presentation():
    # Yeni bir Word belgesi oluştur
    doc = Document()
    
    # Sayfa kenar boşluklarını ayarla
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Başlık ve altlık ekle
    add_header(doc)
    add_footer(doc)
    
    # Kapak Sayfası
    title = doc.add_heading('Hipertansiyon Tahmin Veri Seti\nKeşifsel Veri Analizi ve Makine Öğrenmesi Raporu', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Alt başlık
    subtitle = doc.add_paragraph('Biyoistatistik Dersi Final Ödevi')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    
    # Tarih
    date = doc.add_paragraph(datetime.now().strftime('%d %B %Y'))
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.runs[0].font.size = Pt(12)
    
    doc.add_page_break()
    
    # İçindekiler
    toc = doc.add_heading('İçindekiler', level=1)
    toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contents = [
        '1. Veri Seti Genel Bakış',
        '2. Veri Ön İşleme',
        '3. Keşifsel Veri Analizi (EDA)',
        '4. Makine Öğrenmesi Modeli',
        '5. Model Sonuçları ve Değerlendirme',
        '6. Sonuç ve Öneriler'
    ]
    
    for item in contents:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # Bölüm 1: Veri Seti Genel Bakış
    section1 = doc.add_heading('1. Veri Seti Genel Bakış', level=1)
    section1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    content1 = doc.add_paragraph('''
    Bu çalışmada, Kaggle platformundan alınan hipertansiyon tahmin veri seti kullanılmıştır. 
    Veri seti, çeşitli sağlık göstergeleri ve risk faktörlerini içermektedir.
    ''')
    content1.paragraph_format.line_spacing = 1.5
    
    points1 = [
        'Kaggle\'dan alınan hipertansiyon tahmin veri seti',
        'Çeşitli sağlık göstergeleri ve risk faktörleri',
        'Hedef değişken: Risk (Hipertansiyon Riski)',
        'Toplam 4240 örnek ve 13 özellik'
    ]
    
    for point in points1:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # Bölüm 2: Veri Ön İşleme
    section2 = doc.add_heading('2. Veri Ön İşleme', level=1)
    section2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    content2 = doc.add_paragraph('''
    Veri seti üzerinde aşağıdaki ön işleme adımları gerçekleştirilmiştir:
    ''')
    content2.paragraph_format.line_spacing = 1.5
    
    points2 = [
        'Eksik değerlerin temizlenmesi',
        'Kategorik değişkenlerin sayısallaştırılması',
        'Veri setinin analiz için hazırlanması',
        'Eğitim ve test seti olarak ayrılması (%80-%20)'
    ]
    
    for point in points2:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
    
    doc.add_page_break()
    
    # Bölüm 3: Keşifsel Veri Analizi (EDA)
    section3 = doc.add_heading('3. Keşifsel Veri Analizi (EDA)', level=1)
    section3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Her görselleştirmeyi ekle
    images = [
        ('hipertansiyon_dagilimi.png', 'Hipertansiyon Risk Dağılımı'),
        ('korelasyon_matrisi.png', 'Özellikler Arası Korelasyon Matrisi'),
        ('yas_dagilimi.png', 'Yaş Dağılımı'),
        ('bmi_dagilimi.png', 'Vücut Kitle İndeksi (BMI) Dağılımı'),
        ('kan_basinci_dagilimlari.png', 'Kan Basıncı Dağılımları'),
        ('sayisal_ozellikler_kutu_grafik.png', 'Sayısal Özelliklerin Kutu Grafikleri'),
        ('ozellik_cift_grafik.png', 'Özellikler Arası İlişkiler')
    ]
    
    for image_file, title in images:
        if os.path.exists(image_file):
            doc.add_page_break()
            heading = doc.add_heading(title, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Görsel açıklaması
            desc = doc.add_paragraph(f'Şekil {images.index((image_file, title)) + 1}: {title}')
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc.runs[0].font.italic = True
            
            doc.add_picture(image_file, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Bölüm 4: Makine Öğrenmesi Modeli
    doc.add_page_break()
    section4 = doc.add_heading('4. Makine Öğrenmesi Modeli', level=1)
    section4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    content4 = doc.add_paragraph('''
    Bu çalışmada, hipertansiyon riski tahmini için Lojistik Regresyon modeli kullanılmıştır.
    Lojistik regresyon, ikili sınıflandırma problemleri için yaygın olarak kullanılan 
    bir makine öğrenmesi algoritmasıdır.
    ''')
    content4.paragraph_format.line_spacing = 1.5
    
    ml_points = [
        'Kullanılan Algoritma: Lojistik Regresyon',
        'Veri Bölünmesi: %80 eğitim, %20 test',
        'Hedef Değişken: Hipertansiyon Riski (0: Düşük Risk, 1: Yüksek Risk)',
        'Özellik Sayısı: 12 (Risk hariç)',
        'Model Parametreleri: max_iter=1000'
    ]
    
    for point in ml_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
    
    # Bölüm 5: Model Sonuçları ve Değerlendirme
    doc.add_page_break()
    section5 = doc.add_heading('5. Model Sonuçları ve Değerlendirme', level=1)
    section5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Model sonuçlarını oku
    if os.path.exists('ml_results.txt'):
        with open('ml_results.txt', 'r', encoding='utf-8') as f:
            results = f.read()
        
        results_para = doc.add_paragraph(results)
        results_para.paragraph_format.line_spacing = 1.5
    
    # Model görsellerini ekle
    ml_images = [
        ('ml_confusion_matrix.png', 'Karmaşıklık Matrisi'),
        ('ml_roc_curve.png', 'ROC Eğrisi')
    ]
    
    for image_file, title in ml_images:
        if os.path.exists(image_file):
            doc.add_page_break()
            heading = doc.add_heading(title, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Görsel açıklaması
            desc = doc.add_paragraph(f'Şekil {len(images) + ml_images.index((image_file, title)) + 1}: {title}')
            desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            desc.runs[0].font.italic = True
            
            doc.add_picture(image_file, width=Inches(6))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Bölüm 6: Sonuç ve Öneriler
    doc.add_page_break()
    section6 = doc.add_heading('6. Sonuç ve Öneriler', level=1)
    section6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    content6 = doc.add_paragraph('''
    Bu çalışmada, hipertansiyon riski tahmini için kapsamlı bir veri analizi ve 
    makine öğrenmesi uygulaması gerçekleştirilmiştir.
    ''')
    content6.paragraph_format.line_spacing = 1.5
    
    conclusion_points = [
        'Lojistik regresyon modeli %88 doğruluk oranı ile başarılı sonuçlar vermiştir',
        'ROC-AUC değeri 0.94 ile modelin yüksek ayırt edicilik gücüne sahip olduğunu göstermektedir',
        'Kesinlik (%83) ve duyarlılık (%78) değerleri dengeli bir performans sergilemektedir',
        'Model, hipertansiyon riski tahmini için klinik uygulamalarda kullanılabilir',
        'Gelecek çalışmalarda daha karmaşık algoritmalar (Random Forest, XGBoost) denenebilir'
    ]
    
    for point in conclusion_points:
        p = doc.add_paragraph(point, style='List Bullet')
        p.paragraph_format.line_spacing = 1.5
    
    # Teşekkür Sayfası
    doc.add_page_break()
    thanks = doc.add_heading('Teşekkürler', 0)
    thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Belgeyi kaydet
    doc.save('hipertansiyon_ml_raporu.docx')
    print("Makine öğrenmesi sonuçları ile güncellenmiş Word raporu başarıyla oluşturuldu: hipertansiyon_ml_raporu.docx")

if __name__ == "__main__":
    create_presentation() 